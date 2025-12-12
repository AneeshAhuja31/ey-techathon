"""Document processing service for handling uploads and vectorization."""
import os
import uuid
import logging
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import asyncio

from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models.document import Document
from app.services.vector_store import get_vector_store
from app.core.config import settings
from app.db.database import AsyncSessionLocal

logger = logging.getLogger(__name__)

# Upload directory
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


class DocumentService:
    """Service for processing and managing documents."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.vector_store = get_vector_store()

    async def process_upload(
        self,
        file: UploadFile,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process an uploaded file: save, extract text, chunk, and vectorize.

        Args:
            file: Uploaded file
            user_id: Optional user identifier

        Returns:
            Document information including ID and status
        """
        # Validate file type - only PDF supported for now
        file_ext = self._get_file_extension(file.filename)
        if file_ext != "pdf":
            raise ValueError(f"Only PDF files are supported. Got: {file_ext}")

        # Check file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to start

        max_size = settings.max_file_size_mb * 1024 * 1024
        if file_size > max_size:
            raise ValueError(f"File too large. Maximum size is {settings.max_file_size_mb}MB")

        # Generate unique filename
        doc_id = str(uuid.uuid4())
        stored_filename = f"{doc_id}.{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, stored_filename)

        # Save file
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)

        # Create document record
        document = Document(
            id=doc_id,
            user_id=user_id,
            filename=stored_filename,
            original_filename=file.filename,
            file_type=file_ext,
            file_size=file_size,
            status="processing",
        )
        self.db.add(document)
        await self.db.commit()

        # Process asynchronously
        asyncio.create_task(self._process_document(doc_id, file_path, file_ext, user_id))

        return {
            "id": doc_id,
            "filename": file.filename,
            "file_type": file_ext,
            "file_size": file_size,
            "status": "processing",
            "message": "Document is being processed. Check status for updates."
        }

    async def _process_document(
        self,
        doc_id: str,
        file_path: str,
        file_type: str,
        user_id: Optional[str]
    ):
        """Process document in background: extract, chunk, vectorize."""
        # Use a new session for background processing
        async with AsyncSessionLocal() as db:
            try:
                # Extract text
                text = await self.extract_text(file_path, file_type)

                if not text or len(text.strip()) < 10:
                    await self._update_document_status_with_db(db, doc_id, "failed", "No text content extracted")
                    return

                # Chunk text
                chunks = self.chunk_text(text)

                if not chunks:
                    await self._update_document_status_with_db(db, doc_id, "failed", "Failed to create text chunks")
                    return

                # Vectorize and store
                metadata = {
                    "user_id": user_id,
                    "filename": os.path.basename(file_path),
                    "file_type": file_type,
                }

                chunk_count = await self.vector_store.add_document(
                    doc_id=doc_id,
                    chunks=chunks,
                    metadata=metadata
                )

                # Update document status
                await self._update_document_status_with_db(
                    db,
                    doc_id,
                    "ready",
                    chunk_count=chunk_count
                )

                logger.info(f"Document {doc_id} processed: {chunk_count} chunks")

            except Exception as e:
                logger.error(f"Error processing document {doc_id}: {e}")
                await self._update_document_status_with_db(db, doc_id, "failed", str(e))

    async def _update_document_status(
        self,
        doc_id: str,
        status: str,
        error: str = None,
        chunk_count: int = None
    ):
        """Update document status in database using instance db."""
        await self._update_document_status_with_db(self.db, doc_id, status, error, chunk_count)

    async def _update_document_status_with_db(
        self,
        db: AsyncSession,
        doc_id: str,
        status: str,
        error: str = None,
        chunk_count: int = None
    ):
        """Update document status in database with provided session."""
        try:
            result = await db.execute(
                select(Document).where(Document.id == doc_id)
            )
            document = result.scalar_one_or_none()

            if document:
                document.status = status
                if error:
                    document.error = error
                if chunk_count is not None:
                    document.chunk_count = chunk_count
                document.updated_at = datetime.utcnow()
                await db.commit()
        except Exception as e:
            logger.error(f"Error updating document status: {e}")

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a file.

        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, xlsx, txt)

        Returns:
            Extracted text content
        """
        try:
            if file_type == "pdf":
                return await self._extract_pdf(file_path)
            elif file_type == "docx":
                return await self._extract_docx(file_path)
            elif file_type == "xlsx":
                return await self._extract_xlsx(file_path)
            elif file_type == "txt":
                return await self._extract_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    async def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from XLSX file."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(v.strip() for v in row_values):
                        text_parts.append(" | ".join(row_values))

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            raise

    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def chunk_text(
        self,
        text: str,
        chunk_size: int = None,
        overlap: int = None
    ) -> List[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk
            chunk_size: Maximum chunk size in characters
            overlap: Overlap between chunks

        Returns:
            List of text chunks
        """
        chunk_size = chunk_size or settings.chunk_size
        overlap = overlap or settings.chunk_overlap

        # Clean text
        text = text.strip()
        if not text:
            return []

        # Split into sentences (simple approach)
        sentences = []
        current = ""
        for char in text:
            current += char
            if char in ".!?\n" and len(current) > 10:
                sentences.append(current.strip())
                current = ""
        if current.strip():
            sentences.append(current.strip())

        # Build chunks from sentences
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= chunk_size:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                # Start new chunk with overlap from previous
                if overlap > 0 and current_chunk:
                    overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    async def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get document by ID."""
        result = await self.db.execute(
            select(Document).where(Document.id == doc_id)
        )
        document = result.scalar_one_or_none()
        return document.to_dict() if document else None

    async def list_documents(
        self,
        user_id: Optional[str] = None,
        limit: int = 50
    ) -> List[Dict[str, Any]]:
        """List documents, optionally filtered by user."""
        query = select(Document).order_by(Document.created_at.desc()).limit(limit)

        if user_id:
            query = query.where(Document.user_id == user_id)

        result = await self.db.execute(query)
        documents = result.scalars().all()
        return [doc.to_dict() for doc in documents]

    async def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its vectors."""
        # Get document
        result = await self.db.execute(
            select(Document).where(Document.id == doc_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            return False

        # Delete from vector store
        await self.vector_store.delete_document(doc_id)

        # Delete file
        file_path = os.path.join(UPLOAD_DIR, document.filename)
        if os.path.exists(file_path):
            os.remove(file_path)

        # Delete from database
        await self.db.delete(document)
        await self.db.commit()

        return True

    async def search_documents(
        self,
        query: str,
        user_id: Optional[str] = None,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Search across documents using vector similarity."""
        return await self.vector_store.search(
            query=query,
            user_id=user_id,
            limit=limit
        )

    def _get_file_extension(self, filename: str) -> str:
        """Get lowercase file extension."""
        if not filename or "." not in filename:
            return ""
        return filename.rsplit(".", 1)[-1].lower()
